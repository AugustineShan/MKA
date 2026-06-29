# -*- coding: utf-8 -*-
"""把 /pjbg 产出的研究结论 + forecast 盈利预测全表注入评级报告 .docx 模板。

输入：
- 研究结论.md（pjbg 四件：标题/短期/中长期/投资建议）
- Agent/forecast/derived_metrics_annual.csv（R08 盈利预测全表数据）
- 空白模板 docx（默认 .claude/skills/pjbg/template/评级报告模板.docx，可由 PJBG_TEMPLATE_PATH 覆盖）

注入：
- R01 报告标题值格 ← 标题
- R06 研究结论正文 ← 短期/中长期/投资建议（加粗随 .md：__line__→bold 段，行内 **x**→bold 子串）
- R08 盈利预测区 ← 10 指标×年份窗 子表（复刻前端 overview「财务总结」口径）

鲁棒定位：按节标题文本找单元格（报告标题/研究结论/盈利预测(年度)），不硬编码行号——
换模板只要节标题文本不变就能工作，变了则显式报错停止。
"""
from __future__ import annotations

import argparse
import csv
import io
import os
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.app_config import get_companies_dir, get_researcher_name, rating_report_year_config
from src.company_paths import rating_reports_dir

DEFAULT_TEMPLATE = (
    Path(__file__).resolve().parent.parent
    / ".claude"
    / "skills"
    / "pjbg"
    / "template"
    / "评级报告模板.docx"
)

SECTION_HEADER_RESEARCH = "研究结论"
SECTION_HEADER_FORECAST = "盈利预测(年度)"
LABEL_REPORT_TITLE = "报告标题"
LABEL_RESEARCHER = "研究员"
LABEL_STOCK_CODE = "股票代码"
CAPTION = "单位：百万元；倍数除外"
TABLE_FONT = "微软雅黑"
TABLE_SIZE = 9.0

# (csv 列, 中文标签, 格式, 是否 major 加粗) —— 顺序/标签/格式对齐评级报告模板 R08 体例
# （ROE 在 EPS 前；营业收入/归母净利润标 (百万)；yoy/毛利率/ROE 1dp% 不带 +；PE/PB/EV/EBITDA 1dp 不带 x）
RATING_METRICS = [
    ("revenue", "营业收入(百万)", "int", True),
    ("revenue_yoy", "营收YOY %", "pct", False),
    ("gross_margin", "毛利率", "pct", False),
    ("n_income_attr_p", "归母净利润(百万)", "int", True),
    ("n_income_attr_p_yoy", "归母净利润YOY %", "pct", False),
    ("roe", "ROE %", "pct", False),
    ("eps", "EPS", "dec2", False),
    ("pe", "PE", "mult", False),
    ("pb", "PB", "mult", False),
    ("ev_ebitda", "EV/EBITDA", "mult", False),
]


# ---------- 公司目录解析（轻量，不拉 fitz/openpyxl） ----------
def resolve_company_dir(raw: str | Path) -> Path:
    p = Path(raw)
    if p.exists() and p.is_dir():
        return p
    text = str(raw).strip()
    if not text:
        raise SystemExit("company 参数为空")
    companies = get_companies_dir()
    if re.fullmatch(r"\d{6}(?:\.(?:SZ|SH|BJ))?", text.upper()):
        code = text.upper().split(".")[0]
        cands = sorted(companies.glob(f"*_{code}"))
    else:
        cands = sorted(companies.glob(f"{text}_*"))
    if len(cands) == 1:
        return cands[0]
    if not cands:
        raise SystemExit(f"未找到匹配的公司目录: {text}")
    raise SystemExit(f"匹配到多个公司目录 {text}: {[c.name for c in cands]}")


def _extract_stock_code(company_dir: Path) -> str | None:
    """从公司目录名（如 `绿联科技_301606`）提取 6 位数字代码。"""
    m = re.search(r"\d{6}", company_dir.name)
    return m.group(0) if m else None


# ---------- 研究结论.md 解析 ----------
def _parse_inline_bold(text: str) -> list[tuple[str, bool]]:
    parts = text.split("**")
    segs = [(part, i % 2 == 1) for i, part in enumerate(parts) if part != ""]
    return segs or [(text, False)]


def parse_research_md(md_text: str) -> tuple[str | None, list[list[tuple[str, bool]]]]:
    """返回 (标题, R06 段落序列)。每段 = [(text, bold), ...] run 段。"""
    title: str | None = None
    r06: list[list[tuple[str, bool]]] = []
    current: str | None = None  # 'short' | 'long' | 'invest'
    for raw in md_text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("__") and line.endswith("__") and len(line) >= 4:
            inner = line[2:-2]
            if inner.startswith("短期"):
                current = "short"
                r06.append([(inner, True)])
            elif inner.startswith("中长期"):
                current = "long"
                r06.append([(inner, True)])
            elif inner.startswith("投资建议"):
                current = "invest"
                r06.append([(inner, True)])
            else:
                if title is None:
                    title = inner
                else:
                    raise SystemExit(f"研究结论.md 出现未预期的 __标题块__: {inner}")
            continue
        if current in ("short", "long", "invest"):
            r06.append(_parse_inline_bold(line))
    return title, r06


# ---------- docx 基础工具 ----------
def _set_run_font(run, name: str, size: float, bold: bool | None) -> None:
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), name)


def _sample_font(cell) -> tuple[str, float, bool | None]:
    """从单元格现有 run 采样字体名/字号/加粗（run 文本已被清空但 rPr 还在）。"""
    for p in cell.paragraphs:
        for r in p.runs:
            name = r.font.name
            if name is None:
                rPr = r._element.find(qn("w:rPr"))
                if rPr is not None:
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is not None:
                        name = rFonts.get(qn("w:eastAsia")) or rFonts.get(qn("w:ascii"))
            size = r.font.size.pt if r.font.size else None
            if name:
                return name, (size or 9.0), r.bold
    return TABLE_FONT, TABLE_SIZE, False


def _clear_cell_paragraphs(cell):
    """清空单元格到单个空段落（移除嵌套子表 + 删多余段落 + 清首段 run），返回首段。"""
    tc = cell._tc
    for tbl in tc.findall(qn("w:tbl")):
        tc.remove(tbl)
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    p0 = paras[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    return p0


def _set_cell(cell, text: str, bold: bool, font: str, size: float, align=None) -> None:
    p = _clear_cell_paragraphs(cell)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, font, size, bold)


def _set_table_borders(table) -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)


def _unique_cells(row):
    seen = set()
    cells = []
    for c in row.cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            cells.append(c)
    return cells


def _find_title_value_cell(table):
    """找「报告标题」标签行，返回其值格（该行第二个唯一单元格）。"""
    return _find_label_value_cell(table, LABEL_REPORT_TITLE)


def _find_label_value_cell(table, label: str):
    """找首个单元格为 label 的标签行，返回其值格（该行第二个唯一单元格）。"""
    for row in table.rows:
        cells = _unique_cells(row)
        if cells and cells[0].text.strip() == label:
            if len(cells) >= 2:
                return cells[1]
    return None


def _find_section_content_cell(table, header_text: str):
    """找节标题行，返回下一行（内容行）的首单元格（全行合并格）。"""
    rows = table.rows
    for i, row in enumerate(rows):
        cells = _unique_cells(row)
        if cells and cells[0].text.strip() == header_text:
            if i + 1 < len(rows):
                return rows[i + 1].cells[0]
    return None


# ---------- 注入 ----------
def inject_title(value_cell, title: str, bold: bool | None = True) -> None:
    """写入标题/研究员值格。bold=True 强制加粗（报告标题用）；
    bold=None 沿用单元格采样加粗（研究员格保留模板原有样式）。"""
    name, size, sampled_bold = _sample_font(value_cell)
    if bold is None:
        bold = sampled_bold
    p0 = _clear_cell_paragraphs(value_cell)
    run = p0.add_run(title)
    _set_run_font(run, name, size, bold)


def inject_research(content_cell, r06_paragraphs: list[list[tuple[str, bool]]]) -> None:
    name, size, _ = _sample_font(content_cell)
    p0 = _clear_cell_paragraphs(content_cell)
    first = True
    for segments in r06_paragraphs:
        # big-part 节标题（短期/中长期/投资建议）：非首段前插一个空段落，做空行分隔
        if (not first and len(segments) == 1 and segments[0][1]
                and segments[0][0].startswith(("短期", "中长期", "投资建议"))):
            content_cell.add_paragraph()
        p = p0 if first else content_cell.add_paragraph()
        first = False
        for text, bold in segments:
            run = p.add_run(text)
            _set_run_font(run, name, size, bold)


def _normalize_doc_font(doc, font: str = TABLE_FONT) -> None:
    """全局统一字体：把文档 body 内所有 run 的字体设为 font（保留字号/加粗/颜色）。

    修模板里 font-less run（典型是「投资要点/正文部分 见附件」占位符）回退到
    doc 默认字体（Calibri/宋体）的问题——保证输出报告字体一律微软雅黑。
    在 XML 层遍历所有 w:r（覆盖正文/表格/嵌套表），idempotent。
    """
    for r_el in doc.element.body.iter(qn("w:r")):
        rPr = r_el.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r_el.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), font)


def _format_value(val, fmt: str) -> str:
    if val is None:
        return ""
    if fmt == "int":
        return f"{val:,.0f}"
    if fmt == "pct":
        return f"{val * 100:.1f}%"
    if fmt == "dec2":
        return f"{val:.2f}"
    if fmt == "mult":
        return f"{val:.1f}"
    return str(val)


def _load_forecast(csv_path: Path) -> dict[str, dict[str, float | None]]:
    data: dict[str, dict[str, float | None]] = {}
    with open(csv_path, encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        cols = [m[0] for m in RATING_METRICS]
        for row in reader:
            period = row.get("period") or row.get("﻿period")
            if not period:
                continue
            data[str(period)] = {
                col: (float(row[col]) if row.get(col) not in (None, "") else None)
                for col in cols
            }
    return data


def _year_window() -> list[tuple[int, bool, str]]:
    cfg = rating_report_year_config()
    years: list[tuple[int, bool, str]] = []
    for y in range(cfg["data_start_year"], cfg["data_end_year"] + 1):
        years.append((y, False, str(y)))
    for y in range(cfg["forecast_start_year"], cfg["forecast_end_year"] + 1):
        if not any(y == yy for yy, _, _ in years):
            years.append((y, True, f"{y}E"))
    return years


def _parse_header_year(header: str) -> int | None:
    """列头 → 年份：'2023'→2023，'2026E'→2026，解析失败 None。"""
    h = header.strip().upper().rstrip("E")
    try:
        return int(h)
    except ValueError:
        return None


def _fill_cell_value(cell, text: str, font: str, size: float, bold) -> None:
    """往值格填一个数字 run，保留单元格段落对齐/边框（只填值，不改格式）。

    清掉多余段落与旧 run，按采样字体写新 run；段落对齐属段落属性，不动。
    """
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    p0 = paras[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    run = p0.add_run(text)
    _set_run_font(run, font, size, bold)


# 行标签 → (csv 列, 格式)：与 RATING_METRICS 同源，用于匹配模板自带表
LABEL_TO_METRIC = {mlabel: (metric, fmt) for metric, mlabel, fmt, _ in RATING_METRICS}


def _fill_existing_forecast_table(table, data: dict) -> None:
    """填充模板自带的盈利预测嵌套表：只往值格写数字，行列名/表头/边框/字体不动。

    按表头列文本解析年份、按行标签匹配指标；匹配不上的行列原样保留。
    """
    rows = table.rows
    hdr = _unique_cells(rows[0])
    # 列索引 → 年份
    col_years: dict[int, int] = {}
    for ci in range(1, len(hdr)):
        year = _parse_header_year(hdr[ci].text)
        if year is not None:
            col_years[ci] = year
    # 数据行：按行标签匹配指标，采样该行标签列字体填值
    for row in rows[1:]:
        cells = _unique_cells(row)
        label = cells[0].text.strip()
        if label not in LABEL_TO_METRIC:
            continue
        metric, fmt = LABEL_TO_METRIC[label]
        font, size, bold = _sample_font(cells[0])
        for ci, year in col_years.items():
            if ci >= len(cells):
                continue
            val = data.get(str(year), {}).get(metric)
            _fill_cell_value(cells[ci], _format_value(val, fmt), font, size, bold)


def _build_forecast_table(content_cell, data: dict) -> None:
    """回退路径：模板无自带盈利预测表时新建（旧空白模板兼容）。"""
    years = _year_window()
    p0 = _clear_cell_paragraphs(content_cell)
    run = p0.add_run(CAPTION)
    _set_run_font(run, TABLE_FONT, TABLE_SIZE, False)

    nyears = len(years)
    table = content_cell.add_table(rows=1, cols=1 + nyears)
    _set_table_borders(table)

    hdr = table.rows[0].cells
    _set_cell(hdr[0], "财务指标", True, TABLE_FONT, TABLE_SIZE, WD_ALIGN_PARAGRAPH.LEFT)
    for i, (_, _, label) in enumerate(years):
        _set_cell(hdr[1 + i], label, True, TABLE_FONT, TABLE_SIZE, WD_ALIGN_PARAGRAPH.CENTER)

    for metric, mlabel, fmt, major in RATING_METRICS:
        cells = table.add_row().cells
        _set_cell(cells[0], mlabel, major, TABLE_FONT, TABLE_SIZE, WD_ALIGN_PARAGRAPH.LEFT)
        for i, (year, _, _) in enumerate(years):
            val = data.get(str(year), {}).get(metric)
            _set_cell(cells[1 + i], _format_value(val, fmt), major, TABLE_FONT, TABLE_SIZE, WD_ALIGN_PARAGRAPH.RIGHT)

    content_cell.add_paragraph()


def inject_forecast(content_cell, company_dir: Path) -> None:
    csv_path = company_dir / "Agent" / "forecast" / "derived_metrics_annual.csv"
    if not csv_path.exists():
        raise SystemExit(f"forecast CSV 不存在，先跑 py -m src.forecast: {csv_path}")
    data = _load_forecast(csv_path)

    # 优先填充模板自带的盈利预测嵌套表（只填数字，不改行列名/格式/边框）
    existing = content_cell.tables
    if existing:
        _fill_existing_forecast_table(existing[0], data)
        return
    # 回退：模板无嵌套表 → 新建（旧空白模板兼容）
    _build_forecast_table(content_cell, data)


def _find_latest_md(company_dir: Path) -> Path:
    d = rating_reports_dir(company_dir)
    cands = sorted(d.glob("*研究结论.md"), key=lambda p: p.name, reverse=True)
    if not cands:
        raise SystemExit(f"未找到研究结论.md: {d}（先跑 /pjbg 生成）")
    return cands[0]


def build_rating_report_docx(
    company_dir: Path | str,
    md_path: Path | str | None = None,
    out_path: Path | str | None = None,
    template_path: Path | str | None = None,
) -> Path:
    company_dir = Path(company_dir)
    template = Path(template_path or os.environ.get("PJBG_TEMPLATE_PATH") or DEFAULT_TEMPLATE)
    if not template.exists():
        raise SystemExit(f"模板不存在: {template}")

    md = Path(md_path) if md_path else _find_latest_md(company_dir)
    title, r06 = parse_research_md(md.read_text(encoding="utf-8"))
    if not title:
        raise SystemExit(f"研究结论.md 未解析到标题: {md}")
    if not r06:
        raise SystemExit(f"研究结论.md 未解析到短期/中长期/投资建议: {md}")

    doc = Document(str(template))
    if not doc.tables:
        raise SystemExit(f"模板里没找到表格: {template}")
    table = doc.tables[0]

    title_cell = _find_title_value_cell(table)
    if title_cell is None:
        raise SystemExit("模板未找到「报告标题」标签行——模板不兼容")
    inject_title(title_cell, title)

    researcher = get_researcher_name()
    if researcher:
        researcher_cell = _find_label_value_cell(table, LABEL_RESEARCHER)
        if researcher_cell is not None:
            inject_title(researcher_cell, researcher, bold=None)

    stock_code = _extract_stock_code(company_dir)
    if stock_code:
        code_cell = _find_label_value_cell(table, LABEL_STOCK_CODE)
        if code_cell is not None:
            inject_title(code_cell, stock_code, bold=None)

    research_cell = _find_section_content_cell(table, SECTION_HEADER_RESEARCH)
    if research_cell is None:
        raise SystemExit("模板未找到「研究结论」节标题——模板不兼容")
    inject_research(research_cell, r06)

    forecast_cell = _find_section_content_cell(table, SECTION_HEADER_FORECAST)
    if forecast_cell is None:
        raise SystemExit("模板未找到「盈利预测(年度)」节标题——模板不兼容")
    inject_forecast(forecast_cell, company_dir)

    # 输出名与 .md 同名前缀（研究结论→评级报告），保日期/公司名一致
    if out_path is None:
        stem = md.stem  # 赛维时代-20260626-研究结论
        docx_stem = stem.replace("研究结论", "评级报告")
        if docx_stem == stem:
            docx_stem = f"{stem}-评级报告"
        out_path = md.parent / f"{docx_stem}.docx"
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    _normalize_doc_font(doc)
    doc.save(str(out))
    return out


def main() -> None:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
    ap = argparse.ArgumentParser(description="注入研究结论+盈利预测全表到评级报告 docx 模板")
    ap.add_argument("--ticker", help="公司代码，如 002946.SZ")
    ap.add_argument("--company-dir", help="公司目录路径")
    ap.add_argument("--md", help="研究结论.md 路径（默认自动找最新）")
    ap.add_argument("--out", help="输出 docx 路径（默认与 .md 同目录同名）")
    ap.add_argument("--template", help="空白模板路径（默认 skill 目录 / PJBG_TEMPLATE_PATH）")
    args = ap.parse_args()

    raw = args.company_dir or args.ticker
    if not raw:
        ap.error("需要 --ticker 或 --company-dir")
    company_dir = resolve_company_dir(raw)
    out = build_rating_report_docx(company_dir, md_path=args.md, out_path=args.out, template_path=args.template)
    print(f"OUT: {out}")


if __name__ == "__main__":
    main()
