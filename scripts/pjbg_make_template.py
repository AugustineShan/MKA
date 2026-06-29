# -*- coding: utf-8 -*-
"""从一份填好的评级报告 docx 派生空白模板。

按节标题文本定位（不硬编码行号），清空可变内容、保留节标题/标签/尾部/合并/字体：
- 节标题行（研究结论/盈利预测(年度)/投资要点/正文部分）：保留，标记下一行为内容行。
- 内容行：清空所有段落文本，保留 run 样式（作注入时的字体采样基底）。
- 表头字段行（其余）：清空「非标签、非空」单元格的文本（即公司名/代码/股价/标题值等），保留标签格。

用法：
    py scripts/pjbg_make_template.py --src "<填好的报告.docx>" --out "<空白模板.docx>"
    py scripts/pjbg_make_template.py   # 默认：伊利样例 → .claude/skills/pjbg/template/评级报告模板.docx
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# 节标题（全行合并格的文本）：命中则该行保留，下一行视为内容行清空
SECTION_HEADERS = {"研究结论", "盈利预测(年度)", "投资要点", "正文部分"}

# 内容行写「见附件」的节标题（投资要点/正文部分另附文档，不内嵌报告）
SEE_APPENDIX_SECTIONS = {"投资要点", "正文部分"}
SEE_APPENDIX_TEXT = "见附件"

# 表头字段行里要保留的标签格文本（其余非空格视为「值」清空）
LABELS = {
    "报告标题", "股票代码", "股票评级", "上次评级", "预期收益率(%)",
    "当前股价", "当前市值(亿元)", "目标价", "年初以来涨跌幅(%)",
    "研究员", "报告类型", "报告日期", "年初以来相对行业涨跌幅(%)",
    "是否重大推荐",
}

ROOT = Path(__file__).resolve().parent.parent
DEFAULT_SRC = Path(r"C:\Users\Sheld\Downloads\2\伊利股份评级报告_更新版.docx")
DEFAULT_OUT = ROOT / ".claude" / "skills" / "pjbg" / "template" / "评级报告模板.docx"


def _unique_cells(row):
    """同一行合并格会重复返回同一 _tc；去重，保留顺序。"""
    seen = set()
    cells = []
    for c in row.cells:
        tc_id = id(c._tc)
        if tc_id not in seen:
            seen.add(tc_id)
            cells.append(c)
    return cells


def _clear_cell_text(cell):
    """清空单元格所有 run 的文本，保留 run 及其 rPr（字体/字号/加粗）作样式基底。"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""


def _set_cell_text_keep_style(cell, text: str) -> None:
    """清到单段后写入 text，保留段落/rPr 样式基底（多空段落单元格也收敛为单段）。"""
    # 采样首个 run 的字体样式（若有）作基底
    name = size = None
    bold = None
    for p in cell.paragraphs:
        for r in p.runs:
            name = r.font.name
            size = r.font.size.pt if r.font.size else None
            bold = r.bold
            if name:
                break
        if name:
            break
    # 移除多余段落，仅保留首段
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    p0 = paras[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    p0.add_run(text)


def _remove_nested_tables(cell):
    """移除单元格内嵌套的子表（如盈利预测区里的预填表），仅清内容行。"""
    tc = cell._tc
    for tbl in tc.findall(qn("w:tbl")):
        tc.remove(tbl)


def blankify(src: Path, out: Path) -> None:
    doc = Document(str(src))
    if not doc.tables:
        raise SystemExit(f"模板里没找到表格：{src}")
    table = doc.tables[0]

    prev_section = None  # 上一行命中的节标题文本，None 表示上一行非节标题
    for row in table.rows:
        cells = _unique_cells(row)
        first_text = cells[0].text.strip() if cells else ""

        if first_text in SECTION_HEADERS:
            prev_section = first_text
            continue

        if prev_section is not None:
            # 内容行（研究结论正文 / 盈利预测区 / 投资要点 / 正文）
            for c in cells:
                _remove_nested_tables(c)
                if prev_section in SEE_APPENDIX_SECTIONS:
                    _set_cell_text_keep_style(c, SEE_APPENDIX_TEXT)
                else:
                    _clear_cell_text(c)
            prev_section = None
            continue

        # 表头字段行：清空「非标签、非空」的值格
        for c in cells:
            t = c.text.strip()
            if t and t not in LABELS:
                _clear_cell_text(c)
        prev_section = None

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"空白模板已写出: {out}")


def main() -> None:
    ap = argparse.ArgumentParser(description="从填好的评级报告派生空白模板")
    ap.add_argument("--src", default=str(DEFAULT_SRC), help="填好的报告 docx 路径")
    ap.add_argument("--out", default=str(DEFAULT_OUT), help="空白模板输出路径")
    args = ap.parse_args()
    blankify(Path(args.src), Path(args.out))


if __name__ == "__main__":
    main()
